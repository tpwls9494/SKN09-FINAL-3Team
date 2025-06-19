from django.shortcuts import render, redirect, get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
from django.db import transaction
from django.db.models.functions import Substr, Cast
from django.db.models import IntegerField, OuterRef, Subquery
from django.utils import timezone
from django.http import HttpResponse, JsonResponse, StreamingHttpResponse
from django.conf import settings
from .forms import TemplateForm
from core.models import Template, Draft, User, Evaluation, Team, TeamLog
from django.template.loader import get_template
from weasyprint import HTML, CSS
from docx import Document
from bs4 import BeautifulSoup, NavigableString
from collections import defaultdict
from io import BytesIO
import requests
import os
import json
from .rag_client import rag_client
from .assist_client import assist_client
import logging


def ai_edit(draft_text, prompt):
    # prompt + draft_text 를 합친 뒤 Assist 모델로 보내는 예시
    full_prompt = f"{prompt}\n\n{draft_text}"
    return assist_client.generate_assist_answer(full_prompt)

def ai_evaluate(draft_text):
    # 평가를 위한 프롬프트 템플릿
    eval_prompt = f"이 초안을 평가해줘:\n\n{draft_text}"
    return assist_client.generate_assist_answer(eval_prompt)

def ai_generate_draft(template_data):
    # 지금 쓰던 mock_ai_generate_draft 대신
    # template_data 를 JSON 직렬화해서 Assist 모델에 넘길 수도 있고
    # 또는 미리 만든 마크다운 뼈대(prompt)를 넘기면 됩니다.
    prompt = f"다음 템플릿을 보고 초안을 작성해줘:\n\n{json.dumps(template_data)}"
    return assist_client.generate_assist_answer(prompt, max_new_tokens=4096)

# — ai_edit —
@csrf_exempt
@require_POST
def assist_edit(request):
    data = json.loads(request.body)
    draft_text = data.get("draft_text", "")
    prompt     = data.get("prompt", "")
    if not draft_text or not prompt:
        return JsonResponse({"success": False, "message": "draft_text와 prompt가 필요합니다."}, status=400)
    try:
        edited = ai_edit(draft_text, prompt)
        return JsonResponse({"success": True, "result": edited})
    except Exception as e:
        return JsonResponse({"success": False, "message": str(e)}, status=500)

# — ai_evaluate —
@csrf_exempt
@require_POST
def assist_evaluate(request):
    data = json.loads(request.body)
    draft_text = data.get("draft_text", "")
    if not draft_text:
        return JsonResponse({"success": False, "message": "draft_text가 필요합니다."}, status=400)
    try:
        evaluation = ai_evaluate(draft_text)
        return JsonResponse({"success": True, "result": evaluation})
    except Exception as e:
        return JsonResponse({"success": False, "message": str(e)}, status=500)

# — ai_generate_draft —
@csrf_exempt
@require_POST
def assist_generate(request):
    data = json.loads(request.body)
    template_data = data.get("template_data")
    if not template_data:
        return JsonResponse({"success": False, "message": "template_data가 필요합니다."}, status=400)
    try:
        draft_md = ai_generate_draft(template_data)
        return JsonResponse({"success": True, "result": draft_md})
    except Exception as e:
        return JsonResponse({"success": False, "message": str(e)}, status=500)


# Mock AI 함수 (FastAPI로 교체 예정)
def mock_ai_edit(draft_text, prompt):
    return f"[Edited by AI based on: '{prompt}']\n\n{draft_text}"

def mock_ai_evaluate(draft_text):
    return "🧠 AI 평가 결과: 이 초안은 기술적 진보성이 우수합니다."

def mock_ai_generate_draft(template_data):
    """템플릿 데이터를 기반으로 마크다운 초안 생성"""
    tech_name = template_data.get('tech_name', '혁신적인 기술 시스템')
    newline = "\n"
    
    draft_content = f"""# 발명의 명칭
{tech_name}

## 기술분야
{template_data.get('tech_description', '')}

## 배경기술
{template_data.get('problem_solved', '')}

기존 기술들은 다양한 한계점을 가지고 있었습니다. 특히 효율성, 정확성, 경제성 측면에서 개선이 필요한 상황이었으며, 이러한 문제점들을 해결하기 위한 혁신적인 접근 방법이 요구되었습니다.

## 해결하려는 과제
{template_data.get('problem_solved', '')}

본 발명은 **{tech_name}**을 통해 기존 기술의 한계를 극복하고, 더 나은 솔루션을 제공하는 것을 주요 목표로 합니다.

## 과제의 해결 수단
{template_data.get('tech_differentiation', '')}

본 발명은 다음과 같은 혁신적인 방법론을 통해 기존 문제를 해결합니다:
- 체계적이고 효율적인 접근 방식
- 사용자 중심의 설계 철학
- 확장 가능한 아키텍처 구현

{f"## 활용 분야{newline}{template_data.get('application_field', '')}{newline}" if template_data.get('application_field') else ''}

## 발명의 효과
본 발명을 통해 다음과 같은 효과를 얻을 수 있습니다:

- **성능 향상**: 기존 기술 대비 현저히 향상된 성능 및 효율성
- **경제성 개선**: 비용 효율적인 솔루션 제공
- **사용자 편의성**: 직관적이고 사용하기 쉬운 인터페이스
- **확장성**: 다양한 분야에서의 실용적 활용 가능성
- **안정성**: 기술적 안정성 및 신뢰성 확보

## 발명을 실시하기 위한 구체적인 내용

### 주요 구성 요소
{template_data.get('components_functions', '')}

### 구현 방식
{template_data.get('implementation_example', '')}

본 발명의 주요 구성 요소들이 유기적으로 연동하여 혁신적인 솔루션을 제공합니다.

{f"### 도면의 간단한 설명{newline}{template_data.get('drawing_description', '')}{newline}" if template_data.get('drawing_description') else ''}

## 특허청구범위

**청구항 1**: {tech_name}에 있어서,
상기 기술의 핵심 구성을 포함하여 혁신적인 기능을 제공하는 것을 특징으로 하는 시스템.

**청구항 2**: 제1항에 있어서,
{template_data.get('components_functions', '').split('.')[0] if template_data.get('components_functions') else '효율적인 데이터 처리 및 분석 기능'}을 추가로 포함하는 것을 특징으로 하는 시스템.

**청구항 3**: 제1항 또는 제2항에 있어서,
사용자 친화적인 인터페이스를 통해 직관적인 조작이 가능한 것을 특징으로 하는 시스템.

**청구항 4**: 제1항 내지 제3항 중 어느 한 항에 있어서,
{template_data.get('tech_differentiation', '').split('.')[0] if template_data.get('tech_differentiation') else '보안 기능을 강화하여 안전한 시스템 운영을 보장'}하는 것을 특징으로 하는 시스템.

---

### 출원인 및 발명자 정보
- **출원인**: {template_data.get('application_info', '')}
- **발명자**: {template_data.get('inventor_info', '')}

---
*※ 본 특허청구범위는 특허법 제42조 제2~5항 및 시행규칙 제21조에 따라 작성되었습니다.*
*※ 모든 기술적 내용은 모범명세서 가이드에 따라 체계적으로 기술되어 있습니다.*"""

    return draft_content

def editor_view(request):
    return render(request, 'assist/editor.html')

@csrf_exempt
def insert_patent_report(request):
    if request.method == 'POST':
        data = json.loads(request.body)

        if data['sc_flag'] == 'create':
            # user 테이블 연동
            user = User.objects.get(username=data.get("user_id"))
            user_code = user.__dict__['user_code']
            print("code", user_code)
            try:
                with transaction.atomic():
                    user_id = data.get("user_id")
                    user = User.objects.get(username=user_id)
                    template = Template.objects.create(
                        user_code = user,
                        tech_name=data.get("tech_name"),
                        tech_description=data.get("tech_description"),
                        problem_solved=data.get("problem_solved"),
                        tech_differentiation=data.get("tech_differentation", "tech_differentiation"),
                        application_field=data.get("application_field"),
                        components_functions=data.get("components_functions"),
                        implementation_example=data.get("implementation_example"),
                        drawing_description=data.get("drawing_description"),
                        application_info=data.get("application_info", "applicant_info"),
                        inventor_info=data.get("inventor_info", "inventors"),
                        date=timezone.now(),
                        template_title=data.get("tech_name"),
                        user_code_id=user_code
                    )

                    draft = Draft.objects.create(
                        draft_name = data.get("tech_name"),
                        draft_title = data.get("tech_name"),
                        version = data.get("version"),
                        template_id = template.template_id,
                        create_draft = data.get("create_draft"),
                        create_time = timezone.now()
                    )
                    
                    return JsonResponse({'status': "success", 'message': '저장 완료', 'template_id': template.template_id, 'draft_id': draft.draft_id})
            except Exception as e:
                return JsonResponse({"status": 'error', 'message': str(e)})
        elif data['sc_flag'] == 'update':
            try:
                drafts = Draft.objects.filter(template_id=data['template_id'])

                if not drafts:
                    return JsonResponse({'status': 'error', 'message': 'No drafts found[404]'}, status=404)
                
                latest_draft = sorted(drafts, key=lambda d: int(d.version[1:]), reverse=True)[0]

                draft_obj = {}
                for key, value in latest_draft.__dict__.items():
                    if not key.startswith('_'):  # _state 등 내부 속성 제외
                        draft_obj[key] = value
                
                if str(draft_obj.get("create_draft")) == str(data['create_draft']):
                    return JsonResponse({'status': 'error', 'meassge': "수정 내용이 없습니다."})
                
                print("before::", draft_obj.get("version"))

                version = str(draft_obj.get("version")).replace("v", "").strip()
                print("number::", version)
                latest_version = f"v{int(version) + 1}"

                draft_obj['version'] = latest_version

                print(draft_obj.get("version"))

                draft = Draft.objects.create(
                    draft_name = draft_obj.get("draft_name"), # 기존거 넣어야함..
                    draft_title = data.get("tech_name"),
                    version = draft_obj.get("version"),
                    template_id = draft_obj.get("template_id"),
                    create_draft = data.get("create_draft"),
                    create_time = timezone.now()
                )
                return JsonResponse({'status': "success", 'message': '저장 완료', 'draft_id': draft.draft_id})
            except Exception as e:
                print("?")
                return JsonResponse({"status": 'error', 'message': str(e)})
        
    return JsonResponse({'status': 'error', "message": "Invalid request [400]"})

@csrf_exempt
def insert_evaluation_result(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        tech_title = data.get("tech_title")
        try:
            template = Template.objects.get(template_title=tech_title)

            latest_version = Draft.objects.filter(template_id=template.template_id).annotate(version_number=Cast(Substr('version', 2), IntegerField())).order_by('-version_number').first()
            
            if latest_version is not None:
                draft = Draft.objects.get(template_id=template.template_id, version=latest_version.__dict__.get("version"))
                draft_id = draft.__dict__.get("draft_id")
                
                evaluation_exist = Evaluation.objects.filter(draft_id=draft_id).first()

                if evaluation_exist:
                    JsonResponse({'status': 'success', 'message': '이미 데이터가 존재하므로 따로 동작하지 않습니다.'})
                else:
                    evaluation = Evaluation.objects.create(
                        content = data.get('content'),
                        created_at = timezone.now(),
                        draft_id = draft_id
                    )
                return JsonResponse({'status': 'success', 'message': '평가 정보가 저장되었습니다.'})
            else:
                return JsonResponse({'status': 'error', 'message': 'Draft not found [404]'})
        except Template.DoesNotExist:
            return JsonResponse({'status':'error', 'message': 'Template not found'})
    return JsonResponse({'status':'error', 'message': "Invalid request [400]"})

@csrf_exempt
def select_my_history(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        user_id = data.get('user_id')

        try:
            user = User.objects.get(username=user_id)
            user_code = user.__dict__['user_code']

            template = Template.objects.filter(user_code_id=user_code)
            template_ids = template.values_list('template_id', flat=True)

            drafts = Draft.objects.filter(template_id__in=template_ids).order_by('template_id', 'version')

            grouped = defaultdict(list)
            for draft in drafts:
                grouped[draft.template_id].append({
                    'id': draft.draft_id,
                    'template_id': draft.template_id,
                    'version': draft.version,
                    'draft_main_name': draft.draft_name,
                    'title': draft.draft_title,
                    'content': draft.create_draft
                })
            return JsonResponse(grouped, safe=False)
        except User.DoesNotExist:
            return JsonResponse({'error': 'User not found'}, status=404)
    return JsonResponse({'status':'error', 'message': "Invalid request [400]"})

@csrf_exempt
def select_team_history(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        user_id = data.get('user_id')

        try:
            user = User.objects.get(username=user_id)
            user_code = user.__dict__['user_code']

            teamLog = TeamLog.objects.filter(user_code=user_code)

            team_id_list = []
            user_code_list = []
            team_group_list = []
            for tl in teamLog:
                if tl not in team_id_list:
                    team_id_list.append(tl.team_id_id)

            for ti in team_id_list:
                teamLogUser = TeamLog.objects.filter(team_id_id=ti)
                teamName = Team.objects.get(team_id=ti)

                for tlu in teamLogUser:
                    user_code_list.append(tlu.user_code_id)

                if len(user_code_list) != 0:
                    for ucl in user_code_list:
                        userInfo = User.objects.get(user_code=ucl)
                        templates = Template.objects.filter(user_code_id=ucl)
                        template_ids = templates.values_list('template_id', flat=True)

                        latest_draft_subquery = Draft.objects.filter(template_id=OuterRef('template_id')).order_by('-version')
                        drafts = Draft.objects.filter(template_id__in=template_ids,
                                                      version=Subquery(latest_draft_subquery.values('version')[:1])).order_by('template_id', 'version')
                        
                        team_grouped = defaultdict(list)

                        if len(drafts) == 0:
                            team_grouped[userInfo.username].append({
                                'userId': userInfo.username,
                                'user_nickname': userInfo.user_nickname
                            })
                        
                        for draft in drafts:
                            team_grouped[userInfo.username].append({
                                'id': draft.draft_id,
                                'template_id': draft.template_id,
                                'version': draft.version,
                                'draft_main_name': draft.draft_name,
                                'title': draft.draft_title,
                                'content': draft.create_draft,
                                'userId': userInfo.username,
                                'user_nickname': userInfo.user_nickname
                            })
                        team_group_list.append(team_grouped)

            return JsonResponse(team_group_list, safe=False)
        except TeamLog.DoesNotExist:
            return JsonResponse({'error': 'User not found'}, status=404)
    return JsonResponse({'status':'error', 'message': "Invalid request [400]"})

@csrf_exempt
def update_history_main_title(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        
        Draft.objects.filter(template_id=data.get("template_id")).update(draft_name=data.get("title"))

        Template.objects.filter(template_id=data.get("template_id")).update(template_title=data.get("title"))

        return JsonResponse({'status': 'success', 'message': '수정 완료'})
    return JsonResponse({'status':'error', 'message': "Invalid request [400]"})

@csrf_exempt
def delete_history_main(request):
    if request.method == 'POST':
        data = json.loads(request.body)
        
        template_id = data.get("template_id")
        Draft.objects.filter(template_id=template_id).delete()
        Template.objects.filter(template_id=template_id).delete()

        return JsonResponse({'status': 'success', 'message': '삭제 완료! '})

    return JsonResponse({'status':'error', 'message': "Invalid request [400]"})

def qa_view(request):
    """AI Q&A 페이지"""
    if request.method == 'POST' and request.headers.get('Content-Type') == 'application/json':
        try:
            data = json.loads(request.body)
            question = data.get('question', '')
            
            # Mock AI 답변 생성 (실제로는 AI 모델 호출)
            answer = mock_qa_answer(question)
            
            return JsonResponse({
                'success': True,
                'answer': answer,
                'message': '답변이 생성되었습니다.'
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e),
                'message': '답변 생성 중 오류가 발생했습니다.'
            })
    
    return render(request, 'assist/qa.html')

def mock_qa_answer(question):
    """Q&A 답변 생성 (Mock)"""
    qa_responses = {
        "특허 출원 절차와 필요 서류는 무엇인가요?": """
        <h3>특허 출원 절차</h3>
        <p>특허 출원은 다음과 같은 단계로 진행됩니다:</p>
        <ol>
          <li><strong>발명의 완성</strong> - 실제로 구현 가능한 발명이어야 함</li>
          <li><strong>선행기술조사</strong> - 출원하려는 발명이 이미 공개되었는지 확인</li>
          <li><strong>특허명세서 작성</strong> - 발명의 내용을 상세히 기술</li>
          <li><strong>출원서류 제출</strong> - 특허청에 필요 서류와 함께 출원</li>
          <li><strong>심사청구</strong> - 출원일로부터 3년 이내에 심사 요청</li>
          <li><strong>특허심사</strong> - 특허청 심사관의 신규성, 진보성 등 검토</li>
          <li><strong>특허등록</strong> - 심사 통과 시 특허권 부여</li>
        </ol>
        """,
        
        "선출원주의는 어떤 적용원리는 무엇인가요?": """
        <h3>선출원주의 개념과 적용원리</h3>
        <p><strong>선출원주의</strong>는 동일한 발명에 대해 여러 출원이 있을 경우, 가장 먼저 출원한 자에게 특허권을 부여하는 제도입니다.</p>
        
        <h4>적용 원리</h4>
        <ul>
          <li><strong>출원일 우선</strong> - 동일한 발명에 대해 출원일이 빠른 것이 우선권을 가집니다</li>
          <li><strong>공개원칙</strong> - 출원일을 기준으로 기술 공개를 촉진합니다</li>
          <li><strong>신속성 장려</strong> - 발명 후 빠른 출원을 유도합니다</li>
        </ul>
        """,
    }
    
    return qa_responses.get(question, f"""
    <h3>답변</h3>
    <p>귀하의 질문 "{question}"에 대해 답변드리겠습니다.</p>
    <p>더 구체적인 질문을 해주시면 더 정확한 답변을 제공할 수 있습니다.</p>
    """)

def download_pdf(request, draft_id):
    data = json.loads(request.body)
    draft = get_object_or_404(Draft, draft_id=draft_id)
    
    # html_content = markdown.markdown(draft.create_draft, extensions=['fenced_code', 'codehilite'])
    html_content = data.get("html")
    template = get_template('assist/includes/draft_pdf_template.html')
    html = template.render({'draft': draft, 'create_draft_html': html_content})

    font_path = os.path.join(settings.BASE_DIR, 'assist', 'static', 'assist', 'fonts', 'malgun.ttf')

    css = CSS(string=f'''
        @font-face {{
              font-family: 'MalgunGothic';
              src: url('file:///{font_path.replace(os.sep,"/")}');
        }}
        body {{
            font-family: 'MalgunGothic';
            font-size: 12pt;
            line-height: 1.6;
        }}
        h1, h2, h3, h4 {{
            font-weigth: bold;
            margin-top: 20px;
            margin-bottom: 10px;
        }}
        p {{
            margin: 5px 0;
        }}
        code {{
            font-family: monospace;
            background-color: #f2f2f2;
            padding: 2px 4px;
            border-radius: 4px;
        }}
        pre {{
            background-color: #f5f5f5;
            padding: 10px;
            border-left: 4px solid #ccc;
            overflow-x: auto;
        }}
        ul, ol {{
            margin-left: 20px;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin-top: 10px;
        }}
        th, td {{
            border: 1px solid #ccc;
            padding: 8px;
        }}
        blockquote {{
            border-left: 4px solid #ccc;
            padding-left: 10px;
            color: #666;
            margin: 10px 0;
        }}
    ''')

    pdf_file = HTML(string=html).write_pdf(stylesheets=[css])
    
    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{draft.draft_title}.pdf"'
    return response

def add_paragraph_with_breaks(doc, text):
    if text.strip():
        doc.add_paragraph(text.strip())

def download_docx(request, draft_id):
    data = json.loads(request.body)
    draft = get_object_or_404(Draft, draft_id=draft_id)
    html_content = data.get("html")
    print(html_content)

    soup = BeautifulSoup(html_content, "html.parser")
    doc = Document()
    last_heading = None

    for elem in soup.contents:
        # 제목 처리
        if elem.name in ['h1', 'h2', 'h3']:
            level = int(elem.name[1])
            heading_text = elem.get_text(strip=True)
            doc.add_heading(heading_text, level=level)
            last_heading = heading_text

        # 리스트 처리
        elif elem.name == 'ul':
            for li in elem.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(strip=True), style='ListBullet')

        # 강조 포함 텍스트 (<strong>청구항 1</strong>: 내용)
        elif elem.name == 'strong':
            text = elem.get_text(strip=True)
            if ':' in text:
                title, content = text.split(':', 1)
                doc.add_heading(title.strip(), level=3)
                add_paragraph_with_breaks(doc, content.strip())
            else:
                doc.add_heading(text.strip(), level=3)

        # 텍스트 노드 (본문으로 취급)
        elif isinstance(elem, NavigableString):
            text = elem.strip()
            if text:  # 중복 방지
                add_paragraph_with_breaks(doc, text)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.read(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{draft.draft_title}.docx"'
    return response

def download_hwp(request, draft_id):
    """HWP 다운로드 (현재는 텍스트 파일로 대체)"""
    draft = get_object_or_404(Draft, draft_id=draft_id)
    
    # HWP 라이브러리가 없으므로 텍스트 파일로 대체
    content = f"{draft.__dict__['draft_name']}\n\n{draft.__dict__['create_draft']}"
    
    response = HttpResponse(content, content_type='text/plain; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{draft.__dict__["draft_title"]}.txt"'
    return response
   
# views.py의 QA 관련 부분 (최종 버전)

import logging
import json
from collections import defaultdict
from django.shortcuts import render
from django.http import JsonResponse, StreamingHttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_method, require_POST
from .rag_client import rag_client

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------
#  QA 관련 함수들
# ---------------------------------------------------------------------

@csrf_exempt
@require_http_methods(["GET"])
def qa_page(request):
    """Q&A 페이지 렌더링"""
    return render(request, "assist/qa.html")

@csrf_exempt
@require_http_methods(["POST"])
def qa_ask(request):
    """Qwen3-8B 모델을 사용한 질문 답변 API"""
    try:
        # JSON 요청 확인
        if not request.headers.get("Content-Type", "").startswith("application/json"):
            return JsonResponse(
                {"success": False, "message": "application/json Content-Type이 필요합니다."},
                status=415
            )

        # 요청 데이터 파싱
        try:
            payload = json.loads(request.body)
        except json.JSONDecodeError:
            return JsonResponse(
                {"success": False, "message": "잘못된 JSON 형식입니다."},
                status=400
            )

        question = payload.get("question", "").strip()
        max_new_tokens = payload.get("max_new_tokens", 512)

        if not question:
            return JsonResponse(
                {"success": False, "message": "질문을 입력해 주세요."},
                status=400,
            )

        # rag_client를 통해 QA 요청
        answer_html = rag_client.generate_qa_answer(question, max_new_tokens)

        return JsonResponse(
            {
                "success": True,
                "answer": answer_html,
                "question": question,
                "message": "답변이 생성되었습니다.",
            }
        )

    except Exception as exc:
        logger.error("qa_ask 예외: %s", exc)
        return JsonResponse(
            {"success": False, "message": f"서버 오류: {str(exc)}"},
            status=500
        )

@csrf_exempt
@require_http_methods(["GET"])
def qa_test_connection(request):
    """QA 서버 연결 테스트"""
    try:
        # 기존 health_check 재사용
        health_result = rag_client.health_check()

        if health_result.get("status") == "error":
            return JsonResponse({
                "status": "error",
                "message": f"QA 서버 연결 실패: {health_result.get('message')}",
                "server_url": rag_client.base_url
            }, status=500)

        # 간단한 QA 테스트
        test_answer = rag_client.generate_qa_answer("연결 테스트", 50)

        return JsonResponse({
            "status": "success",
            "message": "QA 서버 연결 성공",
            "server_url": rag_client.base_url,
            "health_check": health_result,
            "test_answer": test_answer[:100] + "..." if len(test_answer) > 100 else test_answer
        })

    except Exception as e:
        logger.error(f"QA 서버 연결 테스트 실패: {str(e)}")
        return JsonResponse({
            "status": "error",
            "message": f"QA 서버 연결 실패: {str(e)}",
            "server_url": rag_client.base_url
        }, status=500)

# ---------------------------------------------------------------------
#  기존 RAG 관련 함수들 (변경 없음)
# ---------------------------------------------------------------------

@csrf_exempt
@require_http_methods(["POST"])
def test_rag_connection(request):
    """RAG 서버 연결 테스트"""
    try:
        health = rag_client.health_check()
        return JsonResponse({
            "status": "success",
            "health_check": health,
            "server_url": rag_client.base_url
        })
    except Exception as e:
        logger.error(f"RAG 연결 테스트 실패: {str(e)}")
        return JsonResponse({
            "status": "error",
            "message": str(e)
        }, status=500)

@csrf_exempt
@require_http_methods(["POST"])
def ask_question(request):
    """질문하기 API (RAG)"""
    if request.method == "GET":
        return JsonResponse({"message": "POST 요청을 사용하세요"})

    try:
        data = json.loads(request.body)
        query = data.get('query', '').strip()

        if not query:
            return JsonResponse({
                "status": "error",
                "message": "질문을 입력해주세요."
            }, status=400)

        # RAG 서버에 질문 요청
        result = rag_client.generate_answer(
            query=query,
            max_new_tokens=data.get('max_new_tokens', 512),
            top_k=data.get('top_k', 5),
            temperature=data.get('temperature', 0.7)
        )

        return JsonResponse(result)

    except json.JSONDecodeError:
        return JsonResponse({
            "status": "error",
            "message": "잘못된 JSON 형식입니다."
        }, status=400)
    except Exception as e:
        logger.error(f"질문 처리 실패: {str(e)}")
        return JsonResponse({
            "status": "error",
            "message": str(e)
        }, status=500)

@csrf_exempt
@require_http_methods(["GET"])
def ask_question_stream(request):
    """스트리밍 질문하기 API (GET 방식)"""
    try:
        # GET 파라미터에서 데이터 추출
        query = request.GET.get('query', '').strip()
        max_new_tokens = int(request.GET.get('max_new_tokens', 512))
        top_k = int(request.GET.get('top_k', 5))
        temperature = float(request.GET.get('temperature', 0.7))

        if not query:
            def error_generator():
                yield f"data: {json.dumps({'type': 'error', 'message': '질문을 입력해주세요'})}\n\n"

            response = StreamingHttpResponse(
                error_generator(),
                content_type='text/event-stream; charset=utf-8'
            )
            response['Cache-Control'] = 'no-cache'
            response['Connection'] = 'keep-alive'
            response['X-Accel-Buffering'] = 'no'
            return response

        # RunPod RAG 서버에 GET 방식으로 스트리밍 요청
        def stream_from_rag():
            try:
                # FastAPI GET 엔드포인트 호출 (URL 파라미터 사용)
                params = {
                    'query': query,
                    'max_new_tokens': max_new_tokens,
                    'top_k': top_k,
                    'temperature': temperature
                }

                response = requests.get(
                    f'{rag_client.base_url}/api/rag/generate-stream',
                    params=params,
                    stream=True,
                    timeout=300  # 5분 타임아웃
                )

                response.raise_for_status()

                # 스트림 데이터를 실시간으로 전달
                for line in response.iter_lines(decode_unicode=True):
                    if line:
                        if line.startswith('data: '):
                            yield f"{line}\n\n"
                        else:
                            yield f"data: {line}\n\n"

                        # 즉시 전송을 위한 플러시
                        import sys
                        sys.stdout.flush()

            except requests.exceptions.RequestException as e:
                logger.error(f"RAG 서버 스트리밍 요청 실패: {str(e)}")
                error_data = json.dumps({
                    'type': 'error',
                    'message': f'RAG 서버 연결 실패: {str(e)}'
                })
                yield f"data: {error_data}\n\n"
            except Exception as e:
                logger.error(f"스트리밍 처리 중 오류: {str(e)}")
                error_data = json.dumps({
                    'type': 'error',
                    'message': f'처리 중 오류: {str(e)}'
                })
                yield f"data: {error_data}\n\n"

        response = StreamingHttpResponse(
            stream_from_rag(),
            content_type='text/event-stream; charset=utf-8'
        )

        # 실시간 스트리밍을 위한 헤더 설정
        response['Cache-Control'] = 'no-cache'
        response['Connection'] = 'keep-alive'
        response['X-Accel-Buffering'] = 'no'  # nginx 버퍼링 방지

        return response

    except ValueError as e:
        logger.error(f"파라미터 오류: {str(e)}")
        def error_generator():
            yield f"data: {json.dumps({'type': 'error', 'message': f'파라미터 오류: {str(e)}'})}\n\n"

        response = StreamingHttpResponse(
            error_generator(),
            content_type='text/event-stream; charset=utf-8'
        )
        response['Cache-Control'] = 'no-cache'
        return response
    except Exception as e:
        logger.error(f"스트리밍 질문 처리 실패: {str(e)}")
        def error_generator():
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"

        response = StreamingHttpResponse(
            error_generator(),
            content_type='text/event-stream; charset=utf-8'
        )
        response['Cache-Control'] = 'no-cache'
        return response

def rag_demo(request):
    """RAG 시스템 데모 페이지"""
    return render(request, 'assist/rag_demo.html')

def rag_demo_stream(request):
    """RAG 시스템 데모 페이지(스트리밍)"""
    return render(request, 'assist/rag_demo_stream.html')

@csrf_exempt
@require_http_methods(["GET"])
def rag_status(request):
    """RAG 시스템 상태 확인"""
    try:
        health = rag_client.health_check()
        return JsonResponse({
            "status": "success",
            "rag_server": health,
            "server_url": rag_client.base_url
        })
    except Exception as e:
        return JsonResponse({
            "status": "error",
            "message": str(e)
        }, status=500)

#===========================================

# assist

import logging
import time
from django.http import JsonResponse
from .assist_client import assist_client

logger = logging.getLogger(__name__)

@csrf_exempt
def assist_ask(request):
    logger.debug("▶ assist_ask called, method=%s, body=%s", request.method, request.body)
    start = time.time()
    try:
        data = json.loads(request.body)
        prompt = data.get("prompt", "")
        max_new_tokens = data.get("max_new_tokens", 32768/2)

        # 실제 모델 호출
        answer = assist_client.generate_assist_answer(prompt, max_new_tokens)
        elapsed = time.time() - start

        logger.debug("◀ assist_ask returning (%.2fs): %s", elapsed, answer[:200])

        return JsonResponse({
            "success": True,
            "answer": answer,
            "prompt": prompt,
            "time": elapsed,
        })
    except Exception as e:
        elapsed = time.time() - start
        logger.exception("❌ assist_ask exception after %.2fs", elapsed)
        return JsonResponse({
            "success": False,
            "error": str(e),
            "time": elapsed,
        }, status=500)
    
def assist_stream(request):
    query = request.GET.get("query", "").strip()
    max_new = request.GET.get("max_new_tokens", "512")

    if not query:
        return StreamingHttpResponse(
            "data: {\"type\":\"error\",\"message\":\"query가 필요합니다.\"}\n\n",
            content_type="text/event-stream; charset=utf-8",
        )

    # RunPod FastAPI SSE URL
    upstream_url = f"{settings.FASTAPI_BASE_URL}/api/qwen/assist-stream"
    params = {
        "query": query,
        "max_new_tokens": max_new
    }

    def event_generator():
        try:
            resp = requests.get(upstream_url, params=params, stream=True, timeout=300)
            resp.raise_for_status()
            for raw in resp.iter_lines(decode_unicode=True):
                if not raw:
                    continue
                # FastAPI 쪽에서 이미 "data: {...}" 형태로 내려온다고 가정
                # 그대로 내려보내기만 하면 됩니다.
                yield raw + "\n\n"
        except Exception as e:
            err = json.dumps({"type": "error", "message": str(e)})
            yield f"data: {err}\n\n"

    response = StreamingHttpResponse(
        event_generator(),
        content_type="text/event-stream; charset=utf-8",
    )
    # 버퍼링 방지, hop-by-hop 헤더 덮어쓰기
    response["Cache-Control"] = "no-cache"
    response["X-Accel-Buffering"] = "no"
    response["Connection"] = "keep-alive"
    return response